"""Renders tailored resumes and cover letters to PDF and DOCX bytes, and
extracts plain text back out of an uploaded master resume file.
"""

from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt
from jinja2 import Template
from pypdf import PdfReader
from weasyprint import HTML

_RESUME_HTML_TEMPLATE = Template(
    """
<html><head><meta charset="utf-8"><style>
  body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1a1a1a; font-size: 11pt; }
  h1 { font-size: 20pt; margin-bottom: 0; }
  h2 { font-size: 12pt; text-transform: uppercase; letter-spacing: 0.05em;
       border-bottom: 1px solid #ccc; margin-top: 18px; padding-bottom: 2px; }
  .role { font-weight: bold; }
  .dates { float: right; color: #555; }
  ul { margin: 4px 0 12px 18px; }
</style></head><body>
  <h1>{{ full_name }}</h1>
  <p>{{ contact_line }}</p>
  <h2>Summary</h2>
  <p>{{ summary }}</p>
  <h2>Skills</h2>
  <p>{{ skills | join(', ') }}</p>
  <h2>Experience</h2>
  {% for job in experience %}
    <p><span class="role">{{ job.title }} — {{ job.company }}</span>
       <span class="dates">{{ job.dates }}</span></p>
    <ul>{% for bullet in job.bullets %}<li>{{ bullet }}</li>{% endfor %}</ul>
  {% endfor %}
  <h2>Education</h2>
  {% for edu in education %}
    <p>{{ edu.degree }}, {{ edu.institution }} ({{ edu.dates }})</p>
  {% endfor %}
  {% if certifications %}
    <h2>Certifications</h2>
    <p>{{ certifications | join(', ') }}</p>
  {% endif %}
</body></html>
"""
)

_COVER_LETTER_HTML_TEMPLATE = Template(
    """
<html><head><meta charset="utf-8"><style>
  body { font-family: 'Helvetica Neue', Arial, sans-serif; color: #1a1a1a;
         font-size: 11pt; line-height: 1.5; }
</style></head><body>
  {% for paragraph in paragraphs %}<p>{{ paragraph }}</p>{% endfor %}
</body></html>
"""
)


def render_resume_pdf(structured_content: dict, full_name: str, contact_line: str) -> bytes:
    html = _RESUME_HTML_TEMPLATE.render(
        full_name=full_name, contact_line=contact_line, **structured_content
    )
    return HTML(string=html).write_pdf()


def render_resume_docx(structured_content: dict, full_name: str, contact_line: str) -> bytes:
    doc = Document()
    doc.add_heading(full_name, level=0)
    doc.add_paragraph(contact_line)

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(structured_content.get("summary", ""))

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(structured_content.get("skills", [])))

    doc.add_heading("Experience", level=1)
    for job in structured_content.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{job.get('title', '')} — {job.get('company', '')}")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(f"  ({job.get('dates', '')})")
        for bullet in job.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Education", level=1)
    for edu in structured_content.get("education", []):
        doc.add_paragraph(
            f"{edu.get('degree', '')}, {edu.get('institution', '')} ({edu.get('dates', '')})"
        )

    certs = structured_content.get("certifications", [])
    if certs:
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(", ".join(certs))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_cover_letter_pdf(content: str) -> bytes:
    paragraphs = [p for p in content.split("\n\n") if p.strip()]
    html = _COVER_LETTER_HTML_TEMPLATE.render(paragraphs=paragraphs)
    return HTML(string=html).write_pdf()


def render_cover_letter_docx(content: str) -> bytes:
    doc = Document()
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def extract_text_from_upload(filename: str, content: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if lower_name.endswith(".docx"):
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    return content.decode("utf-8", errors="ignore")
